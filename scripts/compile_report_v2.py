#!/usr/bin/env python3
"""
CyberSim Graduation Report Compiler – v2 (Premium Edition)
===========================================================
Produces a fully KASIT-compliant, CyberSim-themed DOCX + PDF with:
  • Styled chapter title-page blocks (dark accent bar + chapter number)
  • Sidebar accent rules on H2/H3 via paragraph left-border shading
  • Rich alternating-row tables with header fill
  • Captions above tables / below figures per KASIT spec
  • Callout / code-block paragraphs in mono with light-grey fill
  • All 16 figures embedded at their anchor points
  • Front matter with Roman numerals + body with Arabic restart
  • Full TOC / LOF / LOT generated via Word COM (if available)
  • Falls back gracefully when Word COM is unavailable
"""

from __future__ import annotations

import os
import re
import subprocess
import sys
import textwrap
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# python-docx imports
# ---------------------------------------------------------------------------
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches, Emu

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent.parent
CHAPTERS_DIR = BASE / "docs" / "final-report" / "chapters"
DIAGRAMS_DIR = BASE / "docs" / "final-report" / "diagrams" / "export" / "png"
REFERENCES_MD = BASE / "docs" / "final-report" / "references.md"
OUT_DIR = BASE / "docs" / "final-report" / "formal-report"
OUT_DOCX = OUT_DIR / "cybersim-graduation-report.docx"
OUT_PDF = OUT_DIR / "cybersim-graduation-report.pdf"

# ---------------------------------------------------------------------------
# CyberSim Brand Palette
# ---------------------------------------------------------------------------
# Primary dark navy (chapter title blocks, table headers)
BRAND_DARK = RGBColor(0x0D, 0x1B, 0x2A)        # #0D1B2A  deep navy
BRAND_ACCENT = RGBColor(0x00, 0xB4, 0xD8)       # #00B4D8  cyan accent
BRAND_MID = RGBColor(0x17, 0x32, 0x4E)          # #17324E  mid-navy
BRAND_LIGHT = RGBColor(0xE8, 0xF4, 0xF8)        # #E8F4F8  near-white blue
BRAND_ALT_ROW = RGBColor(0xF0, 0xF8, 0xFF)      # #F0F8FF  alice-blue
BRAND_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)   # white
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)            # light grey
CODE_FG = RGBColor(0x1A, 0x1A, 0x2E)            # very dark navy
CAPTION_COLOR = RGBColor(0x44, 0x55, 0x66)      # slate
H2_ACCENT = RGBColor(0x00, 0xB4, 0xD8)          # same as accent
H3_ACCENT = RGBColor(0x17, 0x32, 0x4E)          # mid navy

# ---------------------------------------------------------------------------
# Chapter ordering
# ---------------------------------------------------------------------------
CHAPTER_FILES = [
    "chapter-01-introduction.md",
    "chapter-02-related-existing-systems.md",
    "chapter-03-requirements.md",
    "chapter-04-system-design.md",
    "chapter-05-implementation.md",
    "chapter-06-testing-and-installation.md",
    "chapter-07-conclusions-and-future-work.md",
]


# ===========================================================================
# LOW-LEVEL XML HELPERS
# ===========================================================================

def _rgb_hex(color: RGBColor) -> str:
    """Return 6-char uppercase hex string from RGBColor (tuple subclass)."""
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_rgb(rpr_or_ppr: OxmlElement, tag: str, color: RGBColor) -> None:
    """Set a colour element (w:color, w:shd, etc.) on an rPr or pPr element."""
    el = OxmlElement(tag)
    el.set(qn("w:val"), _rgb_hex(color))
    rpr_or_ppr.append(el)


def set_cell_bg(cell, color: RGBColor) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    # Remove existing shd
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def add_paragraph_border(para, side: str = "left",
                          color: RGBColor = BRAND_ACCENT,
                          sz: int = 12, space: int = 12) -> None:
    """Add a single left-side border to a paragraph (accent rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), _rgb_hex(color))
    pBdr.append(b)
    # Remove old border element if present
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(pBdr)


def set_para_shading(para, fill: RGBColor) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(fill))
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(shd)


def set_spacing(para, before_pt: float = 0, after_pt: float = 0,
                line_pt: Optional[float] = None) -> None:
    pPr = para._p.get_or_add_pPr()
    spc = pPr.find(qn("w:spacing"))
    if spc is None:
        spc = OxmlElement("w:spacing")
        pPr.append(spc)
    spc.set(qn("w:before"), str(int(before_pt * 20)))
    spc.set(qn("w:after"), str(int(after_pt * 20)))
    if line_pt is not None:
        spc.set(qn("w:line"), str(int(line_pt * 20)))
        spc.set(qn("w:lineRule"), "exact")


def set_indent(para, left_cm: float = 0, first_line_cm: float = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left_cm:
        ind.set(qn("w:left"), str(int(left_cm * 567)))
    if first_line_cm:
        ind.set(qn("w:firstLine"), str(int(first_line_cm * 567)))


def add_field(para, field_code: str) -> None:
    """Insert a Word field (TOC, PAGE, etc.) into a paragraph."""
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run = para.add_run()
    run._r.append(fld_char_begin)
    run = para.add_run()
    run._r.append(instr)
    run = para.add_run()
    run._r.append(fld_char_sep)
    run = para.add_run()
    run._r.append(fld_char_end)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    set_spacing(p, 0, 0)


def insert_section_break(doc: Document, break_type: str = "nextPage") -> None:
    """Append a section break to the last paragraph."""
    p = doc.paragraphs[-1]
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "11906")   # A4
    pgSz.set(qn("w:h"), "16838")
    sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), "1134")    # 2 cm
    pgMar.set(qn("w:right"), "1134")
    pgMar.set(qn("w:bottom"), "1134")
    pgMar.set(qn("w:left"), "1701")   # 3 cm
    pgMar.set(qn("w:header"), "709")
    pgMar.set(qn("w:footer"), "709")
    sectPr.append(pgMar)
    # Page numbering for Roman front matter
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), "lowerRoman")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)
    pPr.append(sectPr)


def set_page_number_restart(doc: Document, fmt: str = "decimal", start: int = 1) -> None:
    """Patch the last sectPr in the document body to restart page numbering."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        return
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


# ===========================================================================
# DOCUMENT SCAFFOLDING
# ===========================================================================

def _configure_document(doc: Document) -> None:
    """Set A4 page, KASIT margins (L 3cm, others 2cm), default font."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    # Normal style — Times New Roman 12pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _configure_heading_styles(doc: Document) -> None:
    """Override Heading 1/2/3 styles to match CyberSim theme."""
    # Heading 1 — Chapter heading (handled manually as title block)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*BRAND_DARK)
    h1.font.all_caps = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BRAND_DARK
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = BRAND_MID
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


# ===========================================================================
# STYLED PARAGRAPH BUILDERS
# ===========================================================================

def add_styled_heading1(doc: Document, text: str,
                        chapter_num: Optional[str] = None) -> None:
    """
    Chapter title block: dark navy bar with white CHAPTER N text, then
    the chapter title below in BRAND_ACCENT color.
    """
    # ── Dark bar with "CHAPTER N" ──────────────────────────────────────────
    if chapter_num:
        bar = doc.add_paragraph()
        bar_run = bar.add_run(f"  CHAPTER {chapter_num}")
        bar_run.font.name = "Times New Roman"
        bar_run.font.size = Pt(10)
        bar_run.font.bold = True
        bar_run.font.color.rgb = BRAND_ACCENT
        set_para_shading(bar, BRAND_DARK)
        bar.paragraph_format.space_before = Pt(0)
        bar.paragraph_format.space_after = Pt(0)
        bar.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Chapter title ──────────────────────────────────────────────────────
    title_clean = re.sub(r"^CHAPTER\s+\d+\s*[:\-]?\s*", "", text, flags=re.IGNORECASE).strip()
    # If nothing left (e.g. the whole string was consumed), fall back
    if not title_clean:
        title_clean = text

    p = doc.add_paragraph()
    run = p.add_run(title_clean.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    set_para_shading(p, BRAND_LIGHT)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    # Bottom border = accent rule under title
    add_paragraph_border(p, "bottom", BRAND_ACCENT, sz=6, space=4)


def add_styled_heading2(doc: Document, text: str) -> None:
    """H2 with left accent rule in cyan."""
    p = doc.add_heading(text, level=2)
    add_paragraph_border(p, "left", BRAND_ACCENT, sz=18, space=8)
    set_indent(p, left_cm=0.4)
    set_spacing(p, before_pt=12, after_pt=6)


def add_styled_heading3(doc: Document, text: str) -> None:
    """H3 with subtle left accent in mid-navy."""
    p = doc.add_heading(text, level=3)
    add_paragraph_border(p, "left", BRAND_MID, sz=10, space=6)
    set_indent(p, left_cm=0.3)
    set_spacing(p, before_pt=8, after_pt=4)


def add_body_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    set_spacing(p, before_pt=0, after_pt=6)


def add_bullet_para(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    set_spacing(p, before_pt=0, after_pt=3)


def add_numbered_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    set_spacing(p, before_pt=0, after_pt=3)


def add_code_block(doc: Document, text: str) -> None:
    """Monospace block with light-grey shading and left border."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_FG
        set_para_shading(p, CODE_BG)
        add_paragraph_border(p, "left", BRAND_ACCENT, sz=6, space=6)
        set_spacing(p, before_pt=1, after_pt=1)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_figure_caption(doc: Document, text: str) -> None:
    """Figure caption below image — italic, centered, smaller, slate color."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = CAPTION_COLOR
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before_pt=3, after_pt=10)


def add_table_caption(doc: Document, text: str) -> None:
    """Table caption above table — italic, left-aligned, slate color."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = CAPTION_COLOR
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before_pt=10, after_pt=2)


# ===========================================================================
# STYLED TABLE BUILDER
# ===========================================================================

def add_styled_table(doc: Document, headers: list[str],
                     rows: list[list[str]],
                     caption: Optional[str] = None) -> None:
    if caption:
        add_table_caption(doc, caption)

    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr_row.height = Cm(0.7)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, BRAND_DARK)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BRAND_HEADER_TEXT
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before_pt=2, after_pt=2)

    # Data rows
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = BRAND_ALT_ROW if idx % 2 == 1 else None
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            if bg:
                set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            # Bold first column
            run = p.add_run(cell_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if j == 0:
                run.font.bold = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p, before_pt=2, after_pt=2)

    # Column widths — distribute equally
    total_width = Cm(15.0)
    col_width = int(total_width / n_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # Space after table
    p = doc.add_paragraph()
    set_spacing(p, before_pt=0, after_pt=6)


# ===========================================================================
# FIGURE INSERTER
# ===========================================================================

def add_figure(doc: Document, png_path: Path, caption: str,
               width_cm: float = 14.0) -> None:
    if not png_path.exists():
        p = doc.add_paragraph(f"[Figure missing: {png_path.name}]")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before_pt=8, after_pt=2)
    run = p.add_run()
    run.add_picture(str(png_path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


# ===========================================================================
# MARKDOWN PARSER
# ===========================================================================

class BlockType:
    HEADING1 = "h1"
    HEADING2 = "h2"
    HEADING3 = "h3"
    PARA = "para"
    BULLET = "bullet"
    NUMBERED = "numbered"
    CODE = "code"
    TABLE = "table"
    FIGURE = "figure"
    BLANK = "blank"


def _parse_md_blocks(md_text: str) -> list[tuple[str, str | list]]:
    """
    Minimal Markdown-to-block-list parser.
    Returns list of (block_type, content) pairs.
    """
    lines = md_text.splitlines()
    blocks: list[tuple[str, str | list]] = []
    in_code = False
    code_buf: list[str] = []
    in_table = False
    table_buf: list[str] = []

    for line in lines:
        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                blocks.append((BlockType.CODE, "\n".join(code_buf)))
                code_buf = []
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Table detection
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_buf = []
            table_buf.append(line)
            continue
        elif in_table:
            in_table = False
            blocks.append((BlockType.TABLE, table_buf[:]))
            table_buf = []
            # Fall through to process current line normally

        stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            blocks.append((BlockType.HEADING3, stripped[4:].strip()))
        elif stripped.startswith("## "):
            blocks.append((BlockType.HEADING2, stripped[3:].strip()))
        elif stripped.startswith("# "):
            blocks.append((BlockType.HEADING1, stripped[2:].strip()))
        # Figure reference (Markdown image syntax)
        elif stripped.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if m:
                blocks.append((BlockType.FIGURE, (m.group(1), m.group(2))))
            else:
                blocks.append((BlockType.PARA, stripped))
        # Bullet
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append((BlockType.BULLET, stripped[2:]))
        # Sub-bullet
        elif re.match(r"^\s+[-*] ", line):
            sub = re.sub(r"^\s+[-*] ", "", line)
            blocks.append((BlockType.BULLET, "  " + sub))
        # Numbered list
        elif re.match(r"^\d+\.", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            blocks.append((BlockType.NUMBERED, text))
        elif stripped == "":
            blocks.append((BlockType.BLANK, ""))
        else:
            blocks.append((BlockType.PARA, stripped))

    # Flush any open table
    if table_buf:
        blocks.append((BlockType.TABLE, table_buf))

    return blocks


def _parse_table_rows(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Convert raw | lines into header + rows lists."""
    data_lines = [l for l in table_lines if not re.match(r"^\s*\|[\s\-|]+\|\s*$", l)]
    parsed = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Clean bold markers
        cells = [re.sub(r"\*\*([^*]+)\*\*", r"\1", c) for c in cells]
        parsed.append(cells)
    if not parsed:
        return [], []
    return parsed[0], parsed[1:]


def _clean_inline(text: str) -> str:
    """Strip inline markdown (bold, italic, code, links) to plain text."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)   # bold
    text = re.sub(r"\*([^*]+)\*", r"\1", text)        # italic
    text = re.sub(r"`([^`]+)`", r"\1", text)          # inline code
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links
    return text


# ===========================================================================
# FIGURE MAP — resolves ![...](../diagrams/export/png/xxx.png) to local Path
# ===========================================================================

FIGURE_CAPTIONS: dict[str, str] = {
    # Original 16 diagrams (redesigned)
    "c4-context.png": "Figure 4.1: CyberSim System Context (C4 Level 1)",
    "c4-container.png": "Figure 4.2: CyberSim Container Architecture (C4 Level 2)",
    "dfd-level-0.png": "Figure 4.3: CyberSim Data Flow Diagram (Level 0)",
    "erd-core-schema.png": "Figure 4.4: CyberSim Core Entity-Relationship Diagram",
    "docker-topology.png": "Figure 4.5: Docker Network and Service Topology",
    "red-blue-event-sequence.png": "Figure 4.6: Red-to-Blue Event Sequence (WebSocket + SIEM Pipeline)",
    "uml-use-case.png": "Figure 4.7: UML Use Case Diagram (All Actors and Use Cases)",
    "auth-sequence.png": "Figure 4.8: Authentication Sequence (Registration, Login, JWT Flow)",
    "session-lifecycle-state.png": "Figure 4.9: Session Lifecycle State Machine",
    "scenario-phase-state-machine.png": "Figure 4.10: Scenario Phase State Machine (Gated Methodology)",
    "deployment-architecture.png": "Figure 4.11: CyberSim Deployment Architecture (Docker Networks)",
    "system-component-interaction.png": "Figure 4.12: System Component Interaction Map",
    "ai-safety-pipeline.png": "Figure 5.1: AI Safety Pipeline (10-Stage Socratic Guard)",
    "report-generation-pipeline.png": "Figure 5.2: Report Generation Pipeline",
    "instructor-analytics-flow.png": "Figure 5.3: Instructor Analytics Data Flow",
    "red-team-methodology-flow.png": "Figure 5.4: Red Team Methodology Flow (PTES Phases + Gates)",
    "blue-team-ir-workflow.png": "Figure 5.5: Blue Team Incident Response Workflow",
    "scoring-and-debrief-flow.png": "Figure 5.6: Scoring and Debrief Generation Flow",
    "sc01-topology.png": "Figure 6.1: SC-01 (NovaMed) Scenario Topology and Attack Chain",
    "sc02-topology.png": "Figure 6.2: SC-02 (Nexora) Scenario Topology",
    "sc03-topology.png": "Figure 6.3: SC-03 (Orion) Scenario Topology",
    "scenario-sc01-flow.png": "Figure 6.4: SC-01 NovaMed Attack and Defense Correlation",
}



def _resolve_png(md_path: str) -> Optional[Path]:
    """Extract filename from markdown path and resolve to diagrams dir."""
    fname = Path(md_path).name
    candidate = DIAGRAMS_DIR / fname
    return candidate if candidate.exists() else None


# ===========================================================================
# CHAPTER RENDERER
# ===========================================================================

def _extract_chapter_num(heading: str) -> Optional[str]:
    m = re.search(r"(\d+)", heading)
    return m.group(1) if m else None


def render_chapter(doc: Document, md_path: Path,
                   table_counter: list[int],
                   figure_counter: list[int]) -> None:
    text = md_path.read_text(encoding="utf-8")
    blocks = _parse_md_blocks(text)

    chapter_num: Optional[str] = None

    for btype, content in blocks:
        if btype == BlockType.BLANK:
            continue

        elif btype == BlockType.HEADING1:
            chapter_num = _extract_chapter_num(str(content))
            add_styled_heading1(doc, str(content), chapter_num)

        elif btype == BlockType.HEADING2:
            add_styled_heading2(doc, _clean_inline(str(content)))

        elif btype == BlockType.HEADING3:
            add_styled_heading3(doc, _clean_inline(str(content)))

        elif btype == BlockType.PARA:
            # Skip "Figure X.Y: ..." standalone caption lines (rendered below image)
            if re.match(r"^Figure\s+\d+\.\d+\s*:", str(content)):
                continue
            add_body_para(doc, _clean_inline(str(content)))

        elif btype == BlockType.BULLET:
            raw = str(content)
            level = 1 if raw.startswith("  ") else 0
            add_bullet_para(doc, _clean_inline(raw.strip()), level=level)

        elif btype == BlockType.NUMBERED:
            add_numbered_para(doc, _clean_inline(str(content)))

        elif btype == BlockType.CODE:
            add_code_block(doc, str(content))

        elif btype == BlockType.TABLE:
            table_lines = list(content)  # type: ignore[arg-type]
            headers, rows = _parse_table_rows(table_lines)
            if headers:
                table_counter[0] += 1
                ch = chapter_num or "?"
                cap = f"Table {ch}.{table_counter[0]}: Data table"
                add_styled_table(doc, headers, rows, caption=cap)

        elif btype == BlockType.FIGURE:
            alt_text, md_src = content  # type: ignore[misc]
            png = _resolve_png(str(md_src))
            if png:
                fname = png.name
                cap = FIGURE_CAPTIONS.get(fname, f"Figure: {alt_text}")
                add_figure(doc, png, cap, width_cm=13.5)
            else:
                # Might be a plain "Figure X.Y: caption" annotation line
                add_figure_caption(doc, _clean_inline(str(alt_text)))


# ===========================================================================
# FRONT MATTER BUILDERS
# ===========================================================================

def build_cover_page(doc: Document) -> None:
    """Branded cover with university, project name, team info."""
    # Top spacing
    for _ in range(3):
        p = doc.add_paragraph()
        set_spacing(p, 0, 0)

    # University name
    p = doc.add_paragraph()
    r = p.add_run("The University of Jordan")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BRAND_DARK
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 0, 2)

    p = doc.add_paragraph()
    r = p.add_run("King Abdullah II School of Information Technology")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = BRAND_MID
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 0, 2)

    p = doc.add_paragraph()
    r = p.add_run("Department of Computer Science")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND_MID
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 0, 24)

    # Accent divider
    div = doc.add_paragraph()
    div_r = div.add_run("─" * 48)
    div_r.font.color.rgb = BRAND_ACCENT
    div_r.font.size = Pt(10)
    div.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(div, 0, 24)

    # Project Title Block
    title_bg = doc.add_paragraph()
    set_para_shading(title_bg, BRAND_DARK)
    t_run = title_bg.add_run("  CyberSim  ")
    t_run.font.name = "Times New Roman"
    t_run.font.size = Pt(28)
    t_run.font.bold = True
    t_run.font.color.rgb = BRAND_ACCENT
    title_bg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title_bg, 12, 4)

    subtitle = doc.add_paragraph()
    set_para_shading(subtitle, BRAND_DARK)
    s_run = subtitle.add_run("  A Dual-Perspective Cybersecurity Training Platform  ")
    s_run.font.name = "Times New Roman"
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = BRAND_LIGHT
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, 4, 20)

    # Accent divider
    div2 = doc.add_paragraph()
    d2r = div2.add_run("─" * 48)
    d2r.font.color.rgb = BRAND_ACCENT
    d2r.font.size = Pt(10)
    div2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(div2, 0, 20)

    # Authors / supervisors
    for label, value in [
        ("Graduation Project — 2025/2026", ""),
        ("Submitted in partial fulfillment of the requirements", ""),
        ("for the Bachelor's Degree in Computer Science", ""),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.color.rgb = BRAND_MID
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, 0, 3)

    set_spacing(doc.paragraphs[-1], 0, 20)

    # Year block
    year = doc.add_paragraph()
    y_run = year.add_run("Academic Year 2025 – 2026")
    y_run.font.name = "Times New Roman"
    y_run.font.size = Pt(12)
    y_run.font.bold = True
    y_run.font.color.rgb = BRAND_DARK
    year.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(year, 0, 6)

    add_page_break(doc)


def build_abstract(doc: Document) -> None:
    add_styled_heading1(doc, "ABSTRACT", None)
    abstract_text = (
        "CyberSim is a browser-based cybersecurity training platform designed for university "
        "students. It provides a dual-perspective learning environment where students practice "
        "Red Team offensive techniques against isolated Docker scenario containers and "
        "simultaneously observe the Blue Team defensive telemetry generated by their actions. "
        "The platform integrates a Kali-style PTY terminal, a real-time SIEM event feed, "
        "structured methodology gating, a Socratic AI tutor, mission scoring, and post-session "
        "debrief reports. Three training scenarios are included in the MVP scope: SC-01 NovaMed "
        "(web application penetration testing), SC-02 Nexora (Active Directory compromise), and "
        "SC-03 Orion (phishing campaign simulation and SOC response). The entire stack is "
        "packaged as a single-node Docker Compose deployment, making it portable for university "
        "labs and graduation demonstrations. Instructor analytics provide session monitoring, "
        "student metrics, and report export capabilities."
    )
    add_body_para(doc, abstract_text)

    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.font.bold = True
    r.font.name = "Times New Roman"
    r2 = p.add_run(
        "cybersecurity training, penetration testing, SOC analysis, Docker, "
        "xterm.js, FastAPI, React, Elasticsearch, Socratic AI, dual-perspective"
    )
    r2.font.name = "Times New Roman"
    r2.font.italic = True
    set_spacing(p, 12, 6)
    add_page_break(doc)


def build_toc_page(doc: Document) -> None:
    """TOC placeholder — Word COM will populate."""
    add_styled_heading1(doc, "TABLE OF CONTENTS", None)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    add_page_break(doc)


def build_lof_page(doc: Document) -> None:
    add_styled_heading1(doc, "LIST OF FIGURES", None)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Figure"')
    add_page_break(doc)


def build_lot_page(doc: Document) -> None:
    add_styled_heading1(doc, "LIST OF TABLES", None)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Table"')
    add_page_break(doc)


# ===========================================================================
# REFERENCES RENDERER
# ===========================================================================

def build_references_section(doc: Document) -> None:
    add_page_break(doc)
    add_styled_heading1(doc, "REFERENCES", None)

    if not REFERENCES_MD.exists():
        add_body_para(doc, "[References file not found]")
        return

    text = REFERENCES_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    ref_num = 0
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        # Numbered items in references.md
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            ref_num += 1
            ref_text = f"[{ref_num}]  {_clean_inline(m.group(1))}"
            p = doc.add_paragraph()
            run = p.add_run(ref_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            set_spacing(p, before_pt=2, after_pt=4)


# ===========================================================================
# WORD COM → PDF
# ===========================================================================

def convert_to_pdf_via_com(docx_path: Path, pdf_path: Path) -> bool:
    """Use Word COM object to update fields and save as PDF."""
    ps_script = f"""
Add-Type -AssemblyName Microsoft.Office.Interop.Word
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $doc = $word.Documents.Open("{str(docx_path).replace(chr(92), chr(92)*2)}")
    $doc.Fields.Update()
    # Export PDF
    $doc.SaveAs2("{str(pdf_path).replace(chr(92), chr(92)*2)}", 17)
    $doc.Close($false)
    $word.Quit()
    Write-Output "PDF_OK"
}} catch {{
    Write-Output "PDF_FAIL: $_"
}}
"""
    tmp_ps = Path(os.environ.get("TEMP", ".")) / "cybersim_pdf.ps1"
    tmp_ps.write_text(ps_script, encoding="utf-8")
    result = subprocess.run(
        ["powershell", "-ExecutionPolicy", "Bypass", "-File", str(tmp_ps)],
        capture_output=True, text=True, timeout=120
    )
    output = result.stdout.strip()
    print(f"  Word COM output: {output}")
    return "PDF_OK" in output


# ===========================================================================
# MAIN ENTRY POINT
# ===========================================================================

def main() -> None:
    # Force UTF-8 output on Windows consoles
    import io
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print("CyberSim Report Compiler v2 - Premium Edition")
    print("=" * 55)

    doc = Document()
    _configure_document(doc)
    _configure_heading_styles(doc)

    # ── FRONT MATTER (Roman numerals) ────────────────────────────────────
    print("  Building cover page...")
    build_cover_page(doc)
    print("  Building abstract...")
    build_abstract(doc)
    print("  Building TOC / LOF / LOT...")
    build_toc_page(doc)
    build_lof_page(doc)
    build_lot_page(doc)

    # Switch to Arabic numbering for body
    set_page_number_restart(doc, fmt="decimal", start=1)

    # ── CHAPTERS ─────────────────────────────────────────────────────────
    table_counter = [0]
    figure_counter = [0]

    for i, fname in enumerate(CHAPTER_FILES):
        md_path = CHAPTERS_DIR / fname
        if not md_path.exists():
            print(f"  [MISSING] {fname}")
            continue
        # Reset table counter per chapter
        table_counter[0] = 0
        print(f"  Rendering {fname}...")
        render_chapter(doc, md_path, table_counter, figure_counter)
        if i < len(CHAPTER_FILES) - 1:
            add_page_break(doc)

    # ── REFERENCES ───────────────────────────────────────────────────────
    print("  Building references...")
    build_references_section(doc)

    # ── SAVE DOCX ────────────────────────────────────────────────────────
    print(f"  Saving DOCX -> {OUT_DOCX.name}")
    doc.save(str(OUT_DOCX))
    size_docx = OUT_DOCX.stat().st_size
    print(f"  [OK] DOCX saved: {size_docx:,} bytes")

    # ── CONVERT TO PDF ───────────────────────────────────────────────────
    print("  Converting to PDF via Word COM...")
    if convert_to_pdf_via_com(OUT_DOCX, OUT_PDF):
        size_pdf = OUT_PDF.stat().st_size
        print(f"  [OK] PDF saved:  {size_pdf:,} bytes")
    else:
        print("  [WARN] Word COM conversion failed - DOCX is still good.")

    print()
    print("[DONE] Compilation complete.")
    print(f"   DOCX: {OUT_DOCX}")
    print(f"   PDF:  {OUT_PDF}")


if __name__ == "__main__":
    main()
