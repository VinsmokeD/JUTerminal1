import os
import re
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), edge_data.get('val', 'single'))
            b.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(b)
    tcPr.append(tcBorders)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_toc_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_lof_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\t "Figure" \\c'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_lot_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\t "Table" \\c'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def parse_markdown_formatting(paragraph, text, normal_font="Times New Roman", is_code=False):
    # Regex to extract bold, italic, code
    # **bold** or *italic* or `code`
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.name = normal_font
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = normal_font
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(token)
            run.font.name = normal_font
            if is_code:
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)

def format_paragraph(p, space_before=0, space_after=6, line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def compile_report(workspace_dir):
    doc = Document()
    
    # Configure styles
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    
    # Configure margins for initial section
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.top_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.different_first_page_header_footer = True
    
    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    p_univ = doc.add_paragraph()
    format_paragraph(p_univ, space_before=20, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_univ = p_univ.add_run("THE UNIVERSITY OF JORDAN\nKING ABDULLAH II SCHOOL OF INFORMATION TECHNOLOGY\nDEPARTMENT OF COMPUTER SCIENCE\n\n\n\n\n")
    run_univ.font.name = 'Times New Roman'
    run_univ.font.size = Pt(14)
    run_univ.bold = True
    
    p_title = doc.add_paragraph()
    format_paragraph(p_title, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_title = p_title.add_run("PARALLAX: A BROWSER-BASED DUAL-PERSPECTIVE CYBERSECURITY TRAINING PLATFORM\n\n\n\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    
    p_degree = doc.add_paragraph()
    format_paragraph(p_degree, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_degree = p_degree.add_run("Prepared by:\nMahmoud's Graduation Project Team\n\n\nSupervised by:\nGraduation Project Supervisor\n\n\n\n\n\n")
    run_degree.font.name = 'Times New Roman'
    run_degree.font.size = Pt(12)
    run_degree.bold = True
    
    p_submit = doc.add_paragraph()
    format_paragraph(p_submit, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_submit = p_submit.add_run("Submitted in partial fulfillment of the requirements for the degree of Bachelor of Science in Computer Science\n\n\nMay 2026")
    run_submit.font.name = 'Times New Roman'
    run_submit.font.size = Pt(11)
    
    # ----------------------------------------------------
    # TITLE PAGE (similar but official degree clearance)
    # ----------------------------------------------------
    doc.add_page_break()
    p_clear = doc.add_paragraph()
    format_paragraph(p_clear, space_before=24, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_clear = p_clear.add_run("GRADUATION PROJECT CLEARANCE COMMITTEE\n\n\nThis graduation project report has been approved and accepted in partial fulfillment of the requirements for the degree of Bachelor of Science in Computer Science.\n\n\n\n\nCommittee Members:\n\n1. ____________________________ (Supervisor)\n\n2. ____________________________ (Examiner)\n\n3. ____________________________ (Examiner)\n\n\n\nDate: May 2026")
    run_clear.font.name = 'Times New Roman'
    run_clear.font.size = Pt(12)
    run_clear.bold = True
    
    # ----------------------------------------------------
    # SECTION 2: FRONT MATTER (ROMAN NUMERALS)
    # ----------------------------------------------------
    section2 = doc.add_section()
    section2.different_first_page_header_footer = False
    section2.left_margin = Cm(3.0)
    section2.top_margin = Cm(2.0)
    section2.right_margin = Cm(2.0)
    section2.bottom_margin = Cm(2.0)
    
    # Set Roman lower numbering for Section 2
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'romanLower')
    section2._sectPr.append(pgNumType)
    
    # Setup footer page numbering for Section 2
    footer = section2.footer
    footer.is_linked_to_previous = False
    p_foot = footer.paragraphs[0]
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run()
    run_foot.font.name = 'Times New Roman'
    run_foot.font.size = Pt(10)
    add_page_number(run_foot)

    # 1. ABSTRACT (ENGLISH)
    p_abs_head = doc.add_paragraph()
    format_paragraph(p_abs_head, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_abs_head = p_abs_head.add_run("ABSTRACT")
    run_abs_head.font.name = 'Times New Roman'
    run_abs_head.font.size = Pt(14)
    run_abs_head.bold = True
    
    p_abs = doc.add_paragraph()
    format_paragraph(p_abs, space_before=6, space_after=12, line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    run_abs_text = p_abs.add_run("Parallax is a dual-perspective cybersecurity training platform designed for university students to bridge the gap between theoretical attack tools and real-time defensive monitoring. Traditional training platforms separate offensive practice from security analysis, making it difficult for students to understand the causal relationship between attacker commands and defender alerts. Parallax addresses this by providing a unified, web-based training environment that runs entirely safely within Docker-isolated scenarios. The platform features two core workspaces: a Red Team space featuring an interactive terminal linked to a Kali Linux container, and a Blue Team space featuring a real-time SIEM alert feed, triage, and forensics containment controls. Educational command-to-detection causality is bridged using a real-time event correlation engine, and Socratic guidance is provided via an integrated rate-limited AI tutor. The system's effectiveness is validated through three high-fidelity scenarios: SC-01 (OWASP Web Vulnerabilities), SC-02 (Active Directory Lateral Movement), and SC-03 (Phishing Campaign Simulation). Automated unit, integration, and performance load tests verify the stability, security, and low latency of the platform under concurrent training conditions.")
    run_abs_text.font.name = 'Times New Roman'
    run_abs_text.font.size = Pt(12)
    
    # 2. ABSTRACT (ARABIC)
    doc.add_page_break()
    p_abs_ar_head = doc.add_paragraph()
    format_paragraph(p_abs_ar_head, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_abs_ar_head = p_abs_ar_head.add_run("Ù…Ù„Ø®Øµ Ø§Ù„Ù…Ø´Ø±ÙˆØ¹")
    run_abs_ar_head.font.name = 'Times New Roman'
    run_abs_ar_head.font.size = Pt(14)
    run_abs_ar_head.bold = True
    
    p_abs_ar = doc.add_paragraph()
    p_abs_ar.paragraph_format.bidi = True # Right to left
    format_paragraph(p_abs_ar, space_before=6, space_after=12, line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run_abs_ar_text = p_abs_ar.add_run("ÙŠÙ‚Ø¯Ù… Ù…Ø´Ø±ÙˆØ¹ Parallax Ù…Ù†ØµØ© ØªØ¯Ø±ÙŠØ¨ÙŠØ© Ø«Ù†Ø§Ø¦ÙŠØ© Ø§Ù„Ù…Ù†Ø¸ÙˆØ± Ù„Ù„Ø£Ù…Ù† Ø§Ù„Ø³ÙŠØ¨Ø±Ø§Ù†ÙŠØŒ Ù…ØµÙ…Ù…Ø© Ø®ØµÙŠØµØ§Ù‹ Ù„Ø·Ù„Ø¨Ø© Ø§Ù„Ø¬Ø§Ù…Ø¹Ø§Øª Ù„Ø³Ø¯ Ø§Ù„ÙØ¬ÙˆØ© Ø¨ÙŠÙ† Ø£Ø¯ÙˆØ§Øª Ø§Ù„Ù‡Ø¬ÙˆÙ… ÙˆØªØ¯Ø§Ø¨ÙŠØ± Ø§Ù„Ù…Ø±Ø§Ù‚Ø¨Ø© ÙˆØ§Ù„Ø¯ÙØ§Ø¹ Ø§Ù„ÙÙˆØ±ÙŠØ©. ØªØ±ÙƒØ² Ù…Ù†ØµØ§Øª Ø§Ù„ØªØ¯Ø±ÙŠØ¨ Ø§Ù„ØªÙ‚Ù„ÙŠØ¯ÙŠØ© Ø¹Ù„Ù‰ Ø£Ø­Ø¯ Ø§Ù„Ø¬Ø§Ù†Ø¨ÙŠÙ† Ø¯ÙˆÙ† Ø§Ù„Ø¢Ø®Ø±ØŒ Ù…Ù…Ø§ ÙŠØµØ¹Ù‘Ø¨ Ø¹Ù„Ù‰ Ø§Ù„Ø·Ù„Ø¨Ø© ÙÙ‡Ù… Ø§Ù„Ø¹Ù„Ø§Ù‚Ø© Ø§Ù„Ø³Ø¨Ø¨ÙŠØ© Ø¨ÙŠÙ† Ù‡Ø¬Ù…Ø§Øª Ø§Ù„Ù…Ø®ØªØ±Ù‚ ÙˆØ§Ù„ØªÙ†Ø¨ÙŠÙ‡Ø§Øª Ø§Ù„Ø£Ù…Ù†ÙŠØ© Ø§Ù„ØªÙŠ ØªØ¸Ù‡Ø± Ù„Ø¯Ù‰ Ø§Ù„Ù…Ø­Ù„Ù„ Ø§Ù„Ø£Ù…Ù†ÙŠ. ØªØ¹Ø§Ù„Ø¬ Ù…Ù†ØµØ© Parallax Ù‡Ø°Ø§ Ø§Ù„Ø®Ù„Ù„ Ù…Ù† Ø®Ù„Ø§Ù„ ØªÙˆÙÙŠØ± Ø¨ÙŠØ¦Ø© ØªØ¯Ø±ÙŠØ¨ÙŠØ© Ù…ÙˆØ­Ø¯Ø© ÙˆÙ‚Ø§Ø¦Ù…Ø© Ø¹Ù„Ù‰ Ø§Ù„Ù…ØªØµÙØ­ ØªØ¹Ù…Ù„ Ø¨Ø£Ù…Ø§Ù† ØªØ§Ù… Ø¯Ø§Ø®Ù„ Ø­Ø§ÙˆÙŠØ§Øª Ù…Ø¹Ø²ÙˆÙ„Ø© (Docker Containers). ØªØªÙ…ÙŠØ² Ø§Ù„Ù…Ù†ØµØ© Ø¨ÙˆØ¬ÙˆØ¯ Ù…Ø³Ø§Ø­ØªÙŠ Ø¹Ù…Ù„ Ø±Ø¦ÙŠØ³ÙŠØªÙŠÙ†: Ù…Ø³Ø§Ø­Ø© Ø§Ù„ÙØ±ÙŠÙ‚ Ø§Ù„Ø£Ø­Ù…Ø± (Ø§Ù„Ù‡Ø¬ÙˆÙ…) Ø§Ù„ØªÙŠ ØªÙˆÙØ± Ù…Ø­Ø§ÙƒØ§Ø© Ù„Ø³Ø·Ø± Ø£ÙˆØ§Ù…Ø± Ù†Ø¸Ø§Ù… ÙƒØ§Ù„ÙŠ Ù„ÙŠÙ†ÙƒØ³ØŒ ÙˆÙ…Ø³Ø§Ø­Ø© Ø§Ù„ÙØ±ÙŠÙ‚ Ø§Ù„Ø£Ø²Ø±Ù‚ (Ø§Ù„Ø¯ÙØ§Ø¹) Ø§Ù„ØªÙŠ ØªØ¹Ø±Ø¶ ØªØºØ°ÙŠØ© Ø­ÙŠØ© Ù„ØªÙ†Ø¨ÙŠÙ‡Ø§Øª Ù†Ø¸Ø§Ù… Ø¥Ø¯Ø§Ø±Ø© Ù…Ø¹Ù„ÙˆÙ…Ø§Øª ÙˆØ£Ø­Ø¯Ø§Ø« Ø§Ù„Ø£Ù…Ù† Ø§Ù„Ø³ÙŠØ¨Ø±Ø§Ù†ÙŠ (SIEM) Ù…Ø¹ Ø£Ø¯ÙˆØ§Øª Ø§Ù„ØªØ­Ù„ÙŠÙ„ ÙˆØ§Ù„Ø§Ø³ØªØ¬Ø§Ø¨Ø© Ù„Ù„Ø­ÙˆØ§Ø¯Ø«. ÙŠØªÙ… Ø±Ø¨Ø· Ø§Ù„Ø£Ø­Ø¯Ø§Ø« ÙˆØ¥Ø±Ø³Ø§Ù„Ù‡Ø§ ÙÙˆØ±ÙŠØ§Ù‹ Ø¹Ø¨Ø± Ù…Ø­Ø±Ùƒ Ø±Ø¨Ø· Ø§Ù„Ø£Ø­Ø¯Ø§Ø« Ø§Ù„ÙÙˆØ±ÙŠØŒ Ø¨ÙŠÙ†Ù…Ø§ ÙŠØªÙˆÙ„Ù‰ Ù…Ø¹Ù„Ù‘Ù… Ø°ÙƒÙŠ Ù…Ø¯Ø¹ÙˆÙ… Ø¨Ø§Ù„Ø°ÙƒØ§Ø¡ Ø§Ù„Ø§ØµØ·Ù†Ø§Ø¹ÙŠ ØªÙ‚Ø¯ÙŠÙ… ØªÙˆØ¬ÙŠÙ‡Ø§Øª ØªÙØ§Ø¹Ù„ÙŠØ© Ø³Ù‚Ø±Ø§Ø·ÙŠØ©. ØªÙ… Ø§Ù„ØªØ­Ù‚Ù‚ Ù…Ù† ÙØ§Ø¹Ù„ÙŠØ© Ø§Ù„Ù…Ù†ØµØ© Ù…Ù† Ø®Ù„Ø§Ù„ Ø«Ù„Ø§Ø«Ø© Ø³ÙŠÙ†Ø§Ø±ÙŠÙˆÙ‡Ø§Øª Ø¹Ø§Ù„ÙŠØ© Ø§Ù„Ø¯Ù‚Ø©: Ø³ÙŠÙ†Ø§Ø±ÙŠÙˆ SC-01 (Ø«ØºØ±Ø§Øª ØªØ·Ø¨ÙŠÙ‚Ø§Øª Ø§Ù„ÙˆÙŠØ¨ OWASP)ØŒ ÙˆØ³ÙŠÙ†Ø§Ø±ÙŠÙˆ SC-02 (Ø§Ø®ØªØ±Ø§Ù‚ Ø®ÙˆØ§Ø¯Ù… Ø§Ù„Ø¯Ù„ÙŠÙ„ Ø§Ù„Ù†Ø´Ø· AD ÙˆÙ„ØªØ±Ø§Ù„ Ù…ÙˆÙÙ…Ù†Øª)ØŒ ÙˆØ³ÙŠÙ†Ø§Ø±ÙŠÙˆ SC-03 (Ù…Ø­Ø§ÙƒØ§Ø© Ù‡Ø¬Ù…Ø§Øª Ø§Ù„ØªØµÙŠØ¯ Ø§Ù„Ø¥Ù„ÙƒØªØ±ÙˆÙ†ÙŠ). Ø£Ø«Ø¨ØªØª Ø§Ø®ØªØ¨Ø§Ø±Ø§Øª Ø§Ù„ÙƒÙØ§Ø¡Ø© ÙˆØ§Ù„ØªØ­Ù…ÙŠÙ„ Ø§Ù„ØªÙ„Ù‚Ø§Ø¦ÙŠØ© Ø§Ø³ØªÙ‚Ø±Ø§Ø± ÙˆØ£Ù…Ø§Ù† Ø§Ù„Ù…Ù†ØµØ© ÙˆØ§Ù†Ø®ÙØ§Ø¶ Ø²Ù…Ù† Ø§Ù„Ø§Ø³ØªØ¬Ø§Ø¨Ø© ÙÙŠ Ø¸Ù„ Ø¸Ø±ÙˆÙ Ø§Ù„ØªØ¯Ø±ÙŠØ¨ Ø§Ù„Ù…ØªØ²Ø§Ù…Ù† Ù„Ù„Ø·Ù„Ø¨Ø©.")
    run_abs_ar_text.font.name = 'Times New Roman'
    run_abs_ar_text.font.size = Pt(12)
    
    # 3. ACKNOWLEDGMENTS
    doc.add_page_break()
    p_ack_head = doc.add_paragraph()
    format_paragraph(p_ack_head, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_ack_head = p_ack_head.add_run("ACKNOWLEDGMENTS")
    run_ack_head.font.name = 'Times New Roman'
    run_ack_head.font.size = Pt(14)
    run_ack_head.bold = True
    
    p_ack = doc.add_paragraph()
    format_paragraph(p_ack, space_before=6, space_after=12, line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    run_ack_text = p_ack.add_run("We express our deepest gratitude to our project supervisor for their guidance, feedback, and encouragement throughout the design and implementation of Parallax. We also thank the examiners and committee members of the Department of Computer Science at the King Abdullah II School of Information Technology for their review and helpful recommendations. Finally, we thank our families and peers for their continuous support during this academic endeavor.")
    run_ack_text.font.name = 'Times New Roman'
    run_ack_text.font.size = Pt(12)
    
    # 4. TABLE OF CONTENTS
    doc.add_page_break()
    p_toc = doc.add_paragraph()
    format_paragraph(p_toc, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_toc = p_toc.add_run("TABLE OF CONTENTS")
    run_toc.font.name = 'Times New Roman'
    run_toc.font.size = Pt(14)
    run_toc.bold = True
    add_toc_field(p_toc)
    
    # 5. LIST OF FIGURES
    doc.add_page_break()
    p_lof = doc.add_paragraph()
    format_paragraph(p_lof, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_lof = p_lof.add_run("LIST OF FIGURES")
    run_lof.font.name = 'Times New Roman'
    run_lof.font.size = Pt(14)
    run_lof.bold = True
    add_lof_field(p_lof)
    
    # 6. LIST OF TABLES
    doc.add_page_break()
    p_lot = doc.add_paragraph()
    format_paragraph(p_lot, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_lot = p_lot.add_run("LIST OF TABLES")
    run_lot.font.name = 'Times New Roman'
    run_lot.font.size = Pt(14)
    run_lot.bold = True
    add_lot_field(p_lot)
    
    # 7. LIST OF ABBREVIATIONS
    doc.add_page_break()
    p_abb_head = doc.add_paragraph()
    format_paragraph(p_abb_head, space_before=12, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run_abb_head = p_abb_head.add_run("LIST OF ABBREVIATIONS")
    run_abb_head.font.name = 'Times New Roman'
    run_abb_head.font.size = Pt(14)
    run_abb_head.bold = True
    
    abbrev_list = [
        ("AD", "Active Directory"),
        ("API", "Application Programming Interface"),
        ("AS-REP", "Active Directory Authentication Service Response"),
        ("CTF", "Capture The Flag"),
        ("DC", "Domain Controller"),
        ("DFD", "Data Flow Diagram"),
        ("DKIM", "DomainKeys Identified Mail"),
        ("DMARC", "Domain-based Message Authentication, Reporting, and Conformance"),
        ("EDR", "Endpoint Detection and Response"),
        ("ERD", "Entity Relationship Diagram"),
        ("FR", "Functional Requirement"),
        ("GPP", "Group Policy Preferences"),
        ("HUD", "Heads-Up Display"),
        ("IDOR", "Insecure Direct Object Reference"),
        ("IR", "Incident Response"),
        ("JWT", "JSON Web Token"),
        ("KASIT", "King Abdullah II School of Information Technology"),
        ("LFI", "Local File Inclusion"),
        ("LMS", "Learning Management System"),
        ("MVP", "Minimum Viable Product"),
        ("NFR", "Non-Functional Requirement"),
        ("ORM", "Object-Relational Mapping"),
        ("OWASP", "Open Web Application Security Project"),
        ("PTY", "Pseudoterminal"),
        ("PWA", "Progressive Web Application"),
        ("RCE", "Remote Code Execution"),
        ("ROE", "Rules of Engagement"),
        ("RPM", "Requests Per Minute"),
        ("SAMBA", "Open-source implementation of SMB/CIFS"),
        ("SC", "Scenario"),
        ("SIEM", "Security Information and Event Management"),
        ("SPN", "Service Principal Name"),
        ("SPF", "Sender Policy Framework"),
        ("SQLi", "SQL Injection"),
        ("SSO", "Single Sign-On"),
        ("UJ", "The University of Jordan"),
        ("UML", "Unified Modeling Language"),
        ("WAF", "Web Application Firewall"),
        ("WCAG", "Web Content Accessibility Guidelines"),
        ("WS", "WebSocket")
    ]
    
    # Create Table of Abbreviations
    abb_table = doc.add_table(rows=len(abbrev_list), cols=2)
    abb_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    for idx, (abbr, name) in enumerate(abbrev_list):
        row = abb_table.rows[idx]
        row.cells[0].paragraphs[0].add_run(abbr).bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[1].paragraphs[0].add_run(name)
        row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        set_cell_margins(row.cells[0], top=60, bottom=60, left=100, right=100)
        set_cell_margins(row.cells[1], top=60, bottom=60, left=100, right=100)
        set_cell_borders(row.cells[0], bottom={'val': 'single', 'sz': 4, 'color': 'D3D3D3'})
        set_cell_borders(row.cells[1], bottom={'val': 'single', 'sz': 4, 'color': 'D3D3D3'})
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)
        
    # ----------------------------------------------------
    # SECTION 3: BODY (ARABIC NUMERALS - RESTART AT 1)
    # ----------------------------------------------------
    section3 = doc.add_section()
    section3.different_first_page_header_footer = False
    section3.left_margin = Cm(3.0)
    section3.top_margin = Cm(2.0)
    section3.right_margin = Cm(2.0)
    section3.bottom_margin = Cm(2.0)
    
    pgNumType3 = OxmlElement('w:pgNumType')
    pgNumType3.set(qn('w:fmt'), 'decimal')
    pgNumType3.set(qn('w:start'), '1') # Restart numbering at 1
    section3._sectPr.append(pgNumType3)
    
    footer3 = section3.footer
    footer3.is_linked_to_previous = False
    p_foot3 = footer3.paragraphs[0]
    p_foot3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot3 = p_foot3.add_run()
    run_foot3.font.name = 'Times New Roman'
    run_foot3.font.size = Pt(10)
    add_page_number(run_foot3)
    
    # ----------------------------------------------------
    # READ AND COMPILE CHAPTERS
    # ----------------------------------------------------
    chapters_dir = os.path.join(workspace_dir, "docs", "final-report", "chapters")
    chapter_files = [
        "chapter-01-introduction.md",
        "chapter-02-related-existing-systems.md",
        "chapter-03-requirements.md",
        "chapter-04-system-design.md",
        "chapter-05-implementation.md",
        "chapter-06-testing-and-installation.md",
        "chapter-07-conclusions-and-future-work.md"
    ]
    
    table_index = 1
    
    for ch_idx, ch_file in enumerate(chapter_files, start=1):
        file_path = os.path.join(chapters_dir, ch_file)
        if not os.path.exists(file_path):
            print(f"Error: Chapter file not found: {file_path}")
            continue
            
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Parse chapter markdown lines
        lines = content.split('\n')
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []
        
        # Check if the first line is the chapter heading. If so, handle it.
        # Skip empty lines.
        doc.add_page_break() if ch_idx > 1 else None
        
        for line in lines:
            stripped = line.strip()
            
            # Code block detection
            if stripped.startswith('```'):
                if in_code_block:
                    # End code block. Write code block to docx
                    p_code = doc.add_paragraph()
                    format_paragraph(p_code, space_before=4, space_after=4, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                    # We can use single cell table for neat styling
                    code_tbl = doc.add_table(rows=1, cols=1)
                    code_tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    cell = code_tbl.rows[0].cells[0]
                    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                    cell.width = Inches(6.0)
                    # Light grey background
                    shading = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading)
                    # Border around code
                    set_cell_borders(cell, top={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                                           bottom={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                                           left={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                                           right={'val': 'single', 'sz': 4, 'color': 'D3D3D3'})
                    
                    code_p = cell.paragraphs[0]
                    format_paragraph(code_p, space_before=0, space_after=0, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                    code_text = "\n".join(code_lines)
                    parse_markdown_formatting(code_p, code_text, is_code=True)
                    
                    in_code_block = False
                    code_lines = []
                else:
                    in_code_block = True
                continue
                
            if in_code_block:
                code_lines.append(line)
                continue
                
            # Table detection
            if stripped.startswith('|'):
                if '---' in line:
                    # Skip markdown table divider line
                    continue
                in_table = True
                table_lines.append(line)
                continue
            elif in_table:
                # End of table detected (blank line or heading or paragraph)
                if table_lines:
                    # Parse and write table to docx
                    parse_and_add_table(doc, table_lines, ch_idx, table_index)
                    table_index += 1
                    table_lines = []
                in_table = False
                
            if stripped == "":
                continue
                
            # Headings
            if stripped.startswith('# '):
                # H1
                heading_text = stripped[2:].upper()
                p_head = doc.add_paragraph()
                format_paragraph(p_head, space_before=18, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                p_head.paragraph_format.keep_with_next = True
                run_h = p_head.add_run(heading_text)
                run_h.font.name = 'Times New Roman'
                run_h.font.size = Pt(14)
                run_h.bold = True
            elif stripped.startswith('## '):
                # H2
                heading_text = stripped[3:]
                p_head = doc.add_paragraph()
                format_paragraph(p_head, space_before=12, space_after=6, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                p_head.paragraph_format.keep_with_next = True
                run_h = p_head.add_run(heading_text)
                run_h.font.name = 'Times New Roman'
                run_h.font.size = Pt(12)
                run_h.bold = True
            elif stripped.startswith('### '):
                # H3
                heading_text = stripped[4:]
                p_head = doc.add_paragraph()
                format_paragraph(p_head, space_before=6, space_after=4, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                p_head.paragraph_format.keep_with_next = True
                run_h = p_head.add_run(heading_text)
                run_h.font.name = 'Times New Roman'
                run_h.font.size = Pt(12)
                run_h.bold = True
                run_h.italic = True
            # Images
            elif stripped.startswith('![') and ']' in stripped and '(' in stripped:
                match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
                if match:
                    caption_text = match.group(1)
                    rel_img_path = match.group(2)
                    
                    # Resolve relative path to absolute
                    # Usually: ../diagrams/... or ../evidence/...
                    rel_img_path = rel_img_path.replace('file:///', '').replace('%20', ' ')
                    if rel_img_path.startswith('../'):
                        abs_img_path = os.path.abspath(os.path.join(chapters_dir, rel_img_path))
                    else:
                        abs_img_path = os.path.abspath(os.path.join(workspace_dir, rel_img_path))
                        
                    if os.path.exists(abs_img_path):
                        p_img = doc.add_paragraph()
                        format_paragraph(p_img, space_before=6, space_after=6, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                        run_img = p_img.add_run()
                        run_img.add_picture(abs_img_path, width=Inches(5.8))
                        
                        p_cap = doc.add_paragraph()
                        format_paragraph(p_cap, space_before=4, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                        # Make figure number bold
                        fig_num_match = re.match(r'^(Figure \d+\.\d+)(:?\s*-\s*|\s*:\s*|\s+)(.*)$', caption_text, re.IGNORECASE)
                        if fig_num_match:
                            num_part = fig_num_match.group(1)
                            desc_part = fig_num_match.group(3)
                            r_num = p_cap.add_run(num_part + ": ")
                            r_num.font.name = 'Times New Roman'
                            r_num.font.size = Pt(10)
                            r_num.bold = True
                            
                            r_desc = p_cap.add_run(desc_part)
                            r_desc.font.name = 'Times New Roman'
                            r_desc.font.size = Pt(10)
                        else:
                            r_cap = p_cap.add_run(caption_text)
                            r_cap.font.name = 'Times New Roman'
                            r_cap.font.size = Pt(10)
                    else:
                        print(f"Warning: Image file not found: {abs_img_path}")
            # Figure captions as text (e.g. Figure 4.1: Parallax system context.)
            elif re.match(r'^Figure \d+\.\d+:', stripped, re.IGNORECASE):
                # The image handler already generates the caption, so skip the text line if it is a duplicate caption
                continue
            # List items
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p_item = doc.add_paragraph(style='List Bullet')
                format_paragraph(p_item, space_before=0, space_after=3, line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                parse_markdown_formatting(p_item, stripped[2:])
            elif re.match(r'^\d+\.\s', stripped):
                match = re.match(r'^(\d+)\.\s(.*)$', stripped)
                p_item = doc.add_paragraph(style='List Number')
                format_paragraph(p_item, space_before=0, space_after=3, line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                parse_markdown_formatting(p_item, match.group(2))
            else:
                # Normal paragraph
                p_norm = doc.add_paragraph()
                format_paragraph(p_norm, space_before=0, space_after=6, line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
                parse_markdown_formatting(p_norm, stripped)

        # Flush any remaining table at end of chapter
        if in_table and table_lines:
            parse_and_add_table(doc, table_lines, ch_idx, table_index)
            table_index += 1
            table_lines = []
            
    # ----------------------------------------------------
    # APPEND REFERENCES
    # ----------------------------------------------------
    ref_path = os.path.join(workspace_dir, "docs", "final-report", "references.md")
    if os.path.exists(ref_path):
        doc.add_page_break()
        p_ref_head = doc.add_paragraph()
        format_paragraph(p_ref_head, space_before=18, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p_ref_head.paragraph_format.keep_with_next = True
        run_ref_h = p_ref_head.add_run("REFERENCES")
        run_ref_h.font.name = 'Times New Roman'
        run_ref_h.font.size = Pt(14)
        run_ref_h.bold = True
        
        with open(ref_path, 'r', encoding='utf-8') as f:
            ref_content = f.read()
            
        ref_lines = ref_content.split('\n')
        for line in ref_lines:
            stripped = line.strip()
            if stripped == "" or stripped.startswith('#'):
                continue
            
            p_ref = doc.add_paragraph()
            # References are single spaced or 1.5, with hanging indent
            format_paragraph(p_ref, space_before=3, space_after=6, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            p_ref.paragraph_format.left_indent = Inches(0.5)
            p_ref.paragraph_format.first_line_indent = Inches(-0.5)
            parse_markdown_formatting(p_ref, stripped)

    # ----------------------------------------------------
    # APPEND APPENDICES
    # ----------------------------------------------------
    appendices = [
        ("APPENDIX A: SECURITY AND SAFETY CASE", "docs/final-report/security-and-safety-case.md"),
        ("APPENDIX B: DEPLOYMENT AND OPERATIONS MANUAL", "docs/final-report/deployment-and-operations-manual.md"),
        ("APPENDIX C: SCENARIO DESIGN DOSSIER", "docs/final-report/scenario-design-dossier.md"),
        ("APPENDIX D: SC-01 NOVAMED SCENARIO DOSSIER", "docs/final-report/scenarios/sc-01-novamed-dossier.md"),
        ("APPENDIX E: SC-02 NEXORA SCENARIO DOSSIER", "docs/final-report/scenarios/sc-02-nexora-dossier.md"),
        ("APPENDIX F: SC-03 ORION SCENARIO DOSSIER", "docs/final-report/scenarios/sc-03-orion-dossier.md"),
        ("APPENDIX G: STUDENT USER MANUAL", "docs/final-report/user-manuals/student-manual.md"),
        ("APPENDIX H: INSTRUCTOR USER MANUAL", "docs/final-report/user-manuals/instructor-manual.md"),
        ("APPENDIX I: MAINTAINER OPERATIONS MANUAL", "docs/final-report/user-manuals/maintainer-operations-manual.md"),
        ("APPENDIX J: TESTING AND VERIFICATION EVIDENCE", "docs/final-report/testing-and-verification-evidence.md"),
        ("APPENDIX K: KNOWN LIMITATIONS AND FUTURE WORK", "docs/final-report/known-limitations-and-future-work.md"),
        ("APPENDIX L: ACCESSIBILITY AND USABILITY NOTES", "docs/final-report/accessibility-and-usability-notes.md"),
        ("APPENDIX M: REQUIREMENTS TRACEABILITY MATRIX", "docs/final-report/requirements-traceability-matrix.md"),
        ("APPENDIX N: API REFERENCE", "docs/final-report/api-reference.md"),
        ("APPENDIX O: DATABASE REFERENCE", "docs/final-report/database-reference.md")
    ]
    
    app_idx = 1
    for title, rel_path in appendices:
        file_path = os.path.join(workspace_dir, rel_path)
        if not os.path.exists(file_path):
            print(f"Warning: Appendix file not found: {file_path}")
            continue
            
        doc.add_page_break()
        p_app_head = doc.add_paragraph()
        format_paragraph(p_app_head, space_before=18, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p_app_head.paragraph_format.keep_with_next = True
        run_app_h = p_app_head.add_run(title)
        run_app_h.font.name = 'Times New Roman'
        run_app_h.font.size = Pt(14)
        run_app_h.bold = True
        
        with open(file_path, 'r', encoding='utf-8') as f:
            app_content = f.read()
            
        app_lines = app_content.split('\n')
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []
        
        for line in app_lines:
            stripped = line.strip()
            
            if stripped.startswith('```'):
                if in_code_block:
                    p_code = doc.add_paragraph()
                    format_paragraph(p_code, space_before=4, space_after=4, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                    code_tbl = doc.add_table(rows=1, cols=1)
                    code_tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    cell = code_tbl.rows[0].cells[0]
                    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                    cell.width = Inches(6.0)
                    shading = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading)
                    set_cell_borders(cell, top={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                                           bottom={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                                           left={'val': 'single', 'sz': 4, 'color': 'D3D3D3'},
                                           right={'val': 'single', 'sz': 4, 'color': 'D3D3D3'})
                    code_p = cell.paragraphs[0]
                    format_paragraph(code_p, space_before=0, space_after=0, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                    code_text = "\n".join(code_lines)
                    parse_markdown_formatting(code_p, code_text, is_code=True)
                    in_code_block = False
                    code_lines = []
                else:
                    in_code_block = True
                continue
                
            if in_code_block:
                code_lines.append(line)
                continue
                
            if stripped.startswith('|'):
                if '---' in line:
                    continue
                in_table = True
                table_lines.append(line)
                continue
            elif in_table:
                if table_lines:
                    parse_and_add_table(doc, table_lines, f"App_{app_idx}", table_index)
                    table_index += 1
                    table_lines = []
                in_table = False
                
            if stripped == "":
                continue
                
            # Headings
            if stripped.startswith('# '):
                # Skip primary header since it's already in the Appendix Title
                continue
            elif stripped.startswith('## '):
                p_head = doc.add_paragraph()
                format_paragraph(p_head, space_before=12, space_after=6, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                p_head.paragraph_format.keep_with_next = True
                run_h = p_head.add_run(stripped[3:])
                run_h.font.name = 'Times New Roman'
                run_h.font.size = Pt(12)
                run_h.bold = True
            elif stripped.startswith('### '):
                p_head = doc.add_paragraph()
                format_paragraph(p_head, space_before=6, space_after=4, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                p_head.paragraph_format.keep_with_next = True
                run_h = p_head.add_run(stripped[4:])
                run_h.font.name = 'Times New Roman'
                run_h.font.size = Pt(12)
                run_h.bold = True
                run_h.italic = True
            elif stripped.startswith('![') and ']' in stripped and '(' in stripped:
                match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
                if match:
                    caption_text = match.group(1)
                    rel_img_path = match.group(2)
                    rel_img_path = rel_img_path.replace('file:///', '').replace('%20', ' ')
                    if rel_img_path.startswith('../'):
                        abs_img_path = os.path.abspath(os.path.join(workspace_dir, "docs", "final-report", rel_path.split('/')[-2], rel_img_path))
                    else:
                        abs_img_path = os.path.abspath(os.path.join(workspace_dir, rel_img_path))
                        
                    if os.path.exists(abs_img_path):
                        p_img = doc.add_paragraph()
                        format_paragraph(p_img, space_before=6, space_after=6, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                        run_img = p_img.add_run()
                        run_img.add_picture(abs_img_path, width=Inches(5.8))
                        
                        p_cap = doc.add_paragraph()
                        format_paragraph(p_cap, space_before=4, space_after=12, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                        r_cap = p_cap.add_run(caption_text)
                        r_cap.font.name = 'Times New Roman'
                        r_cap.font.size = Pt(10)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p_item = doc.add_paragraph(style='List Bullet')
                format_paragraph(p_item, space_before=0, space_after=3, line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                parse_markdown_formatting(p_item, stripped[2:])
            elif re.match(r'^\d+\.\s', stripped):
                match = re.match(r'^(\d+)\.\s(.*)$', stripped)
                p_item = doc.add_paragraph(style='List Number')
                format_paragraph(p_item, space_before=0, space_after=3, line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                parse_markdown_formatting(p_item, match.group(2))
            else:
                p_norm = doc.add_paragraph()
                format_paragraph(p_norm, space_before=0, space_after=6, line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
                parse_markdown_formatting(p_norm, stripped)

        if in_table and table_lines:
            parse_and_add_table(doc, table_lines, f"App_{app_idx}", table_index)
            table_index += 1
            table_lines = []
            
        app_idx += 1

    # Ensure output directory exists
    output_dir = os.path.join(workspace_dir, "docs", "final-report", "formal-report")
    os.makedirs(output_dir, exist_ok=True)
    
    docx_path = os.path.join(output_dir, "parallax-graduation-report.docx")
    doc.save(docx_path)
    print(f"Report compiled successfully to: {docx_path}")

def parse_and_add_table(doc, table_lines, ch_idx, table_index):
    # Parse rows
    parsed_rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        parsed_rows.append(cells)
        
    if not parsed_rows:
        return
        
    num_cols = len(parsed_rows[0])
    num_rows = len(parsed_rows)
    
    # Auto caption above the table
    # Guess table caption based on context, or use standard numbered table label
    p_cap = doc.add_paragraph()
    format_paragraph(p_cap, space_before=12, space_after=4, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    p_cap.paragraph_format.keep_with_next = True
    
    run_num = p_cap.add_run(f"Table {ch_idx}.{table_index}: ")
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(10)
    run_num.bold = True
    
    # Determine table title
    title_text = f"Dataset mapping {table_index}"
    if ch_idx == 1:
        if table_index == 1: title_text = "Project Objectives"
        elif table_index == 2: title_text = "Scenario Focus"
        elif table_index == 3: title_text = "Stakeholder Needs"
    elif ch_idx == 3:
        if table_index == 1: title_text = "Target Users"
        elif table_index == 2: title_text = "Functional Requirements"
        elif table_index == 3: title_text = "Non-Functional Requirements"
        elif table_index == 4: title_text = "Usability and UX Goals"
    elif ch_idx == 4:
        if table_index == 1: title_text = "Scenario Database Entities"
        elif table_index == 2: title_text = "Chapter 4 Figure Set"
    elif ch_idx == 5:
        if table_index == 1: title_text = "Repository Layout"
    elif ch_idx == 6:
        if table_index == 1: title_text = "Service URLs and Access Points"
        
    run_title = p_cap.add_run(title_text)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(10)
    
    # Create Word Table
    w_table = doc.add_table(rows=num_rows, cols=num_cols)
    w_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    # Style rows
    for r_idx, row_data in enumerate(parsed_rows):
        w_row = w_table.rows[r_idx]
        is_header = (r_idx == 0)
        
        for c_idx, cell_value in enumerate(row_data):
            cell = w_row.cells[c_idx]
            cell_p = cell.paragraphs[0]
            format_paragraph(cell_p, space_before=3, space_after=3, line_spacing=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            
            run_c = cell_p.add_run(cell_value)
            run_c.font.name = 'Times New Roman'
            if is_header:
                run_c.bold = True
                
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            
            # Apply border
            border_args = {
                'bottom': {'val': 'single', 'sz': 4, 'color': 'A0A0A0'}
            }
            if is_header:
                border_args['top'] = {'val': 'single', 'sz': 6, 'color': '000000'}
                border_args['bottom'] = {'val': 'single', 'sz': 8, 'color': '000000'}
                # Shading for header
                shading = parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading)
            set_cell_borders(cell, **border_args)
            
    # Set proportional column widths
    for row in w_table.rows:
        total_width = Inches(6.0)
        col_width = total_width / num_cols
        for cell in row.cells:
            cell.width = col_width

if __name__ == "__main__":
    if len(sys.argv) > 1:
        workspace = sys.argv[1]
    else:
        workspace = r"c:\Users\Mahmo\OneDrive\Documents\Mahmoud\Graduation Project\JUTerminal1"
    compile_report(workspace)
