#!/usr/bin/env python3
"""
Parallax Graduation Report Compiler - v3 (Final Edition)
========================================================
Supersedes compile_report_v2.py. Produces the complete, KASIT-compliant,
Parallax-themed DOCX + PDF used as the official final graduation report.

v3 additions over v2:
  * All 22 diagrams embedded (use case in Ch3, behavioral/deployment models in
    Ch4, process models in Ch5, scenario topologies in Ch6).
  * Full front matter: Cover, Declaration, Acknowledgments, Abstract,
    Table of Contents, List of Figures, List of Tables, List of Abbreviations.
  * Complete technical appendices (A-M) rendered from the final-report sources:
    traceability matrix, API reference, database reference, architecture atlas,
    security and safety case, scenario dossier, testing evidence, deployment and
    operations manual, student/instructor/maintainer manuals, accessibility
    notes, and known limitations / future work.
  * Appendix-aware table captioning (Table A.1, Table B.2, ...).
  * Mermaid code fences in source markdown are skipped (figures are pre-rendered).
"""

from __future__ import annotations

import os
import re
import subprocess
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent.parent
FR = BASE / "docs" / "final-report"
CHAPTERS_DIR = FR / "chapters"
DIAGRAMS_DIR = FR / "diagrams" / "export" / "png"
REFERENCES_MD = FR / "references.md"
OUT_DIR = FR / "formal-report"
OUT_DOCX = OUT_DIR / "parallax-graduation-report.docx"
OUT_PDF = OUT_DIR / "parallax-graduation-report.pdf"

# ---------------------------------------------------------------------------
# Parallax Brand Palette
# ---------------------------------------------------------------------------
BRAND_DARK = RGBColor(0x0D, 0x1B, 0x2A)
BRAND_ACCENT = RGBColor(0x00, 0xB4, 0xD8)
BRAND_MID = RGBColor(0x17, 0x32, 0x4E)
BRAND_LIGHT = RGBColor(0xE8, 0xF4, 0xF8)
BRAND_ALT_ROW = RGBColor(0xF0, 0xF8, 0xFF)
BRAND_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)
CODE_FG = RGBColor(0x1A, 0x1A, 0x2E)
CAPTION_COLOR = RGBColor(0x44, 0x55, 0x66)

CHAPTER_FILES = [
    "chapter-01-introduction.md",
    "chapter-02-related-existing-systems.md",
    "chapter-03-requirements.md",
    "chapter-04-system-design.md",
    "chapter-05-implementation.md",
    "chapter-06-testing-and-installation.md",
    "chapter-07-conclusions-and-future-work.md",
]

# (Appendix letter, Title, source path relative to final-report)
APPENDICES = [
    ("A", "Requirements Traceability Matrix", "requirements-traceability-matrix.md"),
    ("B", "System API Reference", "api-reference.md"),
    ("C", "Database Reference", "database-reference.md"),
    ("D", "Technical Architecture Atlas", "technical-architecture-atlas.md"),
    ("E", "Security and Safety Case", "security-and-safety-case.md"),
    ("F", "Scenario Design Dossier", "scenario-design-dossier.md"),
    ("G", "Testing and Verification Evidence", "testing-and-verification-evidence.md"),
    ("H", "Deployment and Operations Manual", "deployment-and-operations-manual.md"),
    ("I", "Student User Manual", "user-manuals/student-manual.md"),
    ("J", "Instructor User Manual", "user-manuals/instructor-manual.md"),
    ("K", "Maintainer Operations Manual", "user-manuals/maintainer-operations-manual.md"),
    ("L", "Accessibility and Usability Notes", "accessibility-and-usability-notes.md"),
    ("M", "Known Limitations and Future Work", "known-limitations-and-future-work.md"),
]

ABBREVIATIONS = [
    ("AD", "Active Directory"),
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("CSF", "Cybersecurity Framework (NIST)"),
    ("DFD", "Data Flow Diagram"),
    ("DOCX", "Office Open XML Word Document"),
    ("ERD", "Entity-Relationship Diagram"),
    ("FR", "Functional Requirement"),
    ("IR", "Incident Response"),
    ("JWT", "JSON Web Token"),
    ("KASIT", "King Abdullah II School of Information Technology"),
    ("MITRE ATT&CK", "MITRE Adversarial Tactics, Techniques, and Common Knowledge"),
    ("MVP", "Minimum Viable Product"),
    ("NFR", "Non-Functional Requirement"),
    ("NIST", "National Institute of Standards and Technology"),
    ("OWASP", "Open Worldwide Application Security Project"),
    ("PTES", "Penetration Testing Execution Standard"),
    ("PTY", "Pseudo-Terminal"),
    ("REST", "Representational State Transfer"),
    ("ROE", "Rules of Engagement"),
    ("SC", "Scenario (e.g. SC-01)"),
    ("SIEM", "Security Information and Event Management"),
    ("SOC", "Security Operations Center"),
    ("SQL", "Structured Query Language"),
    ("SVG", "Scalable Vector Graphics"),
    ("TOC", "Table of Contents"),
    ("UI", "User Interface"),
    ("UML", "Unified Modeling Language"),
    ("UX", "User Experience"),
    ("WAF", "Web Application Firewall"),
    ("WS", "WebSocket"),
    ("WSTG", "Web Security Testing Guide (OWASP)"),
]

# ===========================================================================
# LOW-LEVEL XML HELPERS
# ===========================================================================

def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_bg(cell, color: RGBColor) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(color))
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def add_paragraph_border(para, side: str = "left",
                         color: RGBColor = BRAND_ACCENT,
                         sz: int = 12, space: int = 12) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), _rgb_hex(color))
    pBdr.append(b)
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(pBdr)


def set_para_shading(para, fill: RGBColor) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _rgb_hex(fill))
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(shd)


def set_spacing(para, before_pt: float = 0, after_pt: float = 0,
                line_pt: Optional[float] = None) -> None:
    pPr = para._p.get_or_add_pPr()
    spc = pPr.find(qn("w:spacing"))
    if spc is None:
        spc = OxmlElement("w:spacing")
        pPr.append(spc)
    spc.set(qn("w:before"), str(int(before_pt * 20)))
    spc.set(qn("w:after"), str(int(after_pt * 20)))
    if line_pt is not None:
        spc.set(qn("w:line"), str(int(line_pt * 20)))
        spc.set(qn("w:lineRule"), "exact")


def set_indent(para, left_cm: float = 0, first_line_cm: float = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left_cm:
        ind.set(qn("w:left"), str(int(left_cm * 567)))
    if first_line_cm:
        ind.set(qn("w:firstLine"), str(int(first_line_cm * 567)))


def add_field(para, field_code: str) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    for el in (fld_char_begin, instr, fld_char_sep, fld_char_end):
        run = para.add_run()
        run._r.append(el)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    set_spacing(p, 0, 0)


def set_page_number_restart(doc: Document, fmt: str = "decimal", start: int = 1) -> None:
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        return
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def add_page_number_footer(doc: Document) -> None:
    """Centered 'Page X' footer on the default section."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = CAPTION_COLOR
    add_field(p, "PAGE")


# ===========================================================================
# DOCUMENT SCAFFOLDING
# ===========================================================================

def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _configure_heading_styles(doc: Document) -> None:
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*BRAND_DARK)
    h1.font.all_caps = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = BRAND_DARK
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = BRAND_MID
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


# ===========================================================================
# STYLED PARAGRAPH BUILDERS
# ===========================================================================

def add_styled_heading1(doc: Document, text: str,
                        label: Optional[str] = None,
                        label_word: str = "CHAPTER") -> None:
    if label:
        bar = doc.add_paragraph()
        bar_run = bar.add_run(f"  {label_word} {label}")
        bar_run.font.name = "Times New Roman"
        bar_run.font.size = Pt(10)
        bar_run.font.bold = True
        bar_run.font.color.rgb = BRAND_ACCENT
        set_para_shading(bar, BRAND_DARK)
        bar.paragraph_format.space_before = Pt(0)
        bar.paragraph_format.space_after = Pt(0)
        bar.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    title_clean = re.sub(r"^(CHAPTER|APPENDIX)\s+[\w\d]+\s*[:\-]?\s*", "", text,
                         flags=re.IGNORECASE).strip()
    if not title_clean:
        title_clean = text

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(title_clean.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    set_para_shading(p, BRAND_LIGHT)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.page_break_before = False
    add_paragraph_border(p, "bottom", BRAND_ACCENT, sz=6, space=4)


def add_frontmatter_heading(doc: Document, text: str) -> None:
    """Front-matter section heading (not in TOC numbering, styled like H1)."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    set_para_shading(p, BRAND_LIGHT)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    add_paragraph_border(p, "bottom", BRAND_ACCENT, sz=6, space=4)


def add_styled_heading2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    add_paragraph_border(p, "left", BRAND_ACCENT, sz=18, space=8)
    set_indent(p, left_cm=0.4)
    set_spacing(p, before_pt=12, after_pt=6)


def add_styled_heading3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    add_paragraph_border(p, "left", BRAND_MID, sz=10, space=6)
    set_indent(p, left_cm=0.3)
    set_spacing(p, before_pt=8, after_pt=4)


def add_body_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    set_spacing(p, before_pt=0, after_pt=6)


def add_bullet_para(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    set_spacing(p, before_pt=0, after_pt=3)


def add_numbered_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    set_spacing(p, before_pt=0, after_pt=3)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_FG
        set_para_shading(p, CODE_BG)
        add_paragraph_border(p, "left", BRAND_ACCENT, sz=6, space=6)
        set_spacing(p, before_pt=1, after_pt=1)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_figure_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = CAPTION_COLOR
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before_pt=3, after_pt=10)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = CAPTION_COLOR
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before_pt=10, after_pt=2)


# ===========================================================================
# STYLED TABLE BUILDER
# ===========================================================================

def add_styled_table(doc: Document, headers: list[str],
                     rows: list[list[str]],
                     caption: Optional[str] = None) -> None:
    if caption:
        add_table_caption(doc, caption)

    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr_row.height = Cm(0.7)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, BRAND_DARK)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BRAND_HEADER_TEXT
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before_pt=2, after_pt=2)

    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = BRAND_ALT_ROW if idx % 2 == 1 else None
        for j in range(n_cols):
            cell_text = row_data[j] if j < len(row_data) else ""
            cell = row.cells[j]
            if bg:
                set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            if j == 0:
                run.font.bold = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(p, before_pt=2, after_pt=2)

    total_width = Cm(15.5)
    col_width = int(total_width / n_cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    p = doc.add_paragraph()
    set_spacing(p, before_pt=0, after_pt=6)


# ===========================================================================
# FIGURE INSERTER
# ===========================================================================

def add_figure(doc: Document, png_path: Path, caption: str,
               width_cm: float = 13.5) -> None:
    if not png_path.exists():
        p = doc.add_paragraph(f"[Figure missing: {png_path.name}]")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before_pt=8, after_pt=2)
    run = p.add_run()
    run.add_picture(str(png_path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


# ===========================================================================
# FIGURE CAPTION MAP  (filename -> caption; numbering matches appearance order)
# ===========================================================================

FIGURE_CAPTIONS: dict[str, str] = {
    "uml-use-case.png": "Figure 3.1: Parallax Use Case Model (Actors and Use Cases)",
    "c4-context.png": "Figure 4.1: Parallax System Context (C4 Level 1)",
    "c4-container.png": "Figure 4.2: Parallax Container Architecture (C4 Level 2)",
    "dfd-level-0.png": "Figure 4.3: Parallax Data Flow Diagram (Level 0)",
    "erd-core-schema.png": "Figure 4.4: Parallax Core Entity-Relationship Diagram",
    "docker-topology.png": "Figure 4.5: Docker Network and Service Topology",
    "red-blue-event-sequence.png": "Figure 4.6: Red-to-Blue Event Sequence (WebSocket + SIEM Pipeline)",
    "auth-sequence.png": "Figure 4.7: Authentication Sequence (Registration, Login, JWT Flow)",
    "session-lifecycle-state.png": "Figure 4.8: Session Lifecycle State Machine",
    "scenario-phase-state-machine.png": "Figure 4.9: Scenario Phase State Machine (Gated Methodology)",
    "deployment-architecture.png": "Figure 4.10: Parallax Deployment Architecture (Docker Networks)",
    "system-component-interaction.png": "Figure 4.11: System Component Interaction Map",
    "red-team-methodology-flow.png": "Figure 5.1: Red Team Methodology Flow (PTES Phases and Gates)",
    "blue-team-ir-workflow.png": "Figure 5.2: Blue Team Incident Response Workflow",
    "ai-safety-pipeline.png": "Figure 5.3: AI Tutor Safety Pipeline (Socratic Guard)",
    "scoring-and-debrief-flow.png": "Figure 5.4: Scoring and Debrief Generation Flow",
    "report-generation-pipeline.png": "Figure 5.5: Report Generation Pipeline",
    "instructor-analytics-flow.png": "Figure 5.6: Instructor Analytics Data Flow",
    "sc01-topology.png": "Figure 6.1: SC-01 (NovaMed) Scenario Topology",
    "sc02-topology.png": "Figure 6.2: SC-02 (Nexora) Scenario Topology",
    "sc03-topology.png": "Figure 6.3: SC-03 (Orion) Scenario Topology",
    "scenario-sc01-flow.png": "Figure 6.4: SC-01 NovaMed Attack and Defense Correlation",
}


def _resolve_png(md_path: str) -> Optional[Path]:
    fname = Path(md_path).name
    candidate = DIAGRAMS_DIR / fname
    return candidate if candidate.exists() else None


# ===========================================================================
# MARKDOWN PARSER
# ===========================================================================

class BlockType:
    HEADING1 = "h1"
    HEADING2 = "h2"
    HEADING3 = "h3"
    PARA = "para"
    BULLET = "bullet"
    NUMBERED = "numbered"
    CODE = "code"
    TABLE = "table"
    FIGURE = "figure"
    BLANK = "blank"


def _parse_md_blocks(md_text: str) -> list[tuple[str, object]]:
    lines = md_text.splitlines()
    blocks: list[tuple[str, object]] = []
    in_code = False
    code_lang = ""
    code_buf: list[str] = []
    in_table = False
    table_buf: list[str] = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                # Skip mermaid blocks (figures are pre-rendered separately)
                if code_lang.lower() != "mermaid":
                    blocks.append((BlockType.CODE, "\n".join(code_buf)))
                code_buf = []
                code_lang = ""
            else:
                in_code = True
                code_lang = line.strip()[3:].strip()
            continue
        if in_code:
            code_buf.append(line)
            continue

        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_buf = []
            table_buf.append(line)
            continue
        elif in_table:
            in_table = False
            blocks.append((BlockType.TABLE, table_buf[:]))
            table_buf = []

        stripped = line.strip()

        if stripped.startswith("### "):
            blocks.append((BlockType.HEADING3, stripped[4:].strip()))
        elif stripped.startswith("## "):
            blocks.append((BlockType.HEADING2, stripped[3:].strip()))
        elif stripped.startswith("# "):
            blocks.append((BlockType.HEADING1, stripped[2:].strip()))
        elif stripped.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if m:
                blocks.append((BlockType.FIGURE, (m.group(1), m.group(2))))
            else:
                blocks.append((BlockType.PARA, stripped))
        elif stripped.startswith("> "):
            blocks.append((BlockType.PARA, stripped[2:]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append((BlockType.BULLET, stripped[2:]))
        elif re.match(r"^\s+[-*] ", line):
            sub = re.sub(r"^\s+[-*] ", "", line)
            blocks.append((BlockType.BULLET, "  " + sub))
        elif re.match(r"^\d+\.", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            blocks.append((BlockType.NUMBERED, text))
        elif stripped == "":
            blocks.append((BlockType.BLANK, ""))
        else:
            blocks.append((BlockType.PARA, stripped))

    if table_buf:
        blocks.append((BlockType.TABLE, table_buf))
    return blocks


def _parse_table_rows(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    data_lines = [l for l in table_lines if not re.match(r"^\s*\|[\s\-:|]+\|\s*$", l)]
    parsed = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        cells = [re.sub(r"\*\*([^*]+)\*\*", r"\1", c) for c in cells]
        cells = [re.sub(r"`([^`]+)`", r"\1", c) for c in cells]
        parsed.append(cells)
    if not parsed:
        return [], []
    return parsed[0], parsed[1:]


def _clean_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


# ===========================================================================
# GENERIC MARKDOWN RENDERER
# ===========================================================================

def render_markdown(doc: Document, md_path: Path, table_counter: list[int],
                    *, caption_label: str, skip_first_h1: bool = True) -> None:
    text = md_path.read_text(encoding="utf-8")
    blocks = _parse_md_blocks(text)
    seen_h1 = False

    for btype, content in blocks:
        if btype == BlockType.BLANK:
            continue
        elif btype == BlockType.HEADING1:
            if skip_first_h1 and not seen_h1:
                seen_h1 = True
                continue
            seen_h1 = True
            add_styled_heading2(doc, _clean_inline(str(content)))
        elif btype == BlockType.HEADING2:
            add_styled_heading2(doc, _clean_inline(str(content)))
        elif btype == BlockType.HEADING3:
            add_styled_heading3(doc, _clean_inline(str(content)))
        elif btype == BlockType.PARA:
            if re.match(r"^Figure\s+[\w\d]+\.\d+\s*:", str(content)):
                continue
            add_body_para(doc, _clean_inline(str(content)))
        elif btype == BlockType.BULLET:
            raw = str(content)
            level = 1 if raw.startswith("  ") else 0
            add_bullet_para(doc, _clean_inline(raw.strip()), level=level)
        elif btype == BlockType.NUMBERED:
            add_numbered_para(doc, _clean_inline(str(content)))
        elif btype == BlockType.CODE:
            add_code_block(doc, str(content))
        elif btype == BlockType.TABLE:
            headers, rows = _parse_table_rows(list(content))  # type: ignore[arg-type]
            if headers:
                table_counter[0] += 1
                cap = f"Table {caption_label}.{table_counter[0]}"
                add_styled_table(doc, headers, rows, caption=cap)
        elif btype == BlockType.FIGURE:
            alt_text, md_src = content  # type: ignore[misc]
            png = _resolve_png(str(md_src))
            if png:
                cap = FIGURE_CAPTIONS.get(png.name, f"Figure: {_clean_inline(str(alt_text))}")
                add_figure(doc, png, cap)
            else:
                add_figure_caption(doc, _clean_inline(str(alt_text)))


def _extract_chapter_num(heading: str) -> Optional[str]:
    m = re.search(r"(\d+)", heading)
    return m.group(1) if m else None


def render_chapter(doc: Document, md_path: Path, table_counter: list[int]) -> None:
    text = md_path.read_text(encoding="utf-8")
    blocks = _parse_md_blocks(text)
    chapter_num: Optional[str] = None

    for btype, content in blocks:
        if btype == BlockType.BLANK:
            continue
        elif btype == BlockType.HEADING1:
            chapter_num = _extract_chapter_num(str(content))
            add_styled_heading1(doc, str(content), chapter_num, label_word="CHAPTER")
        elif btype == BlockType.HEADING2:
            add_styled_heading2(doc, _clean_inline(str(content)))
        elif btype == BlockType.HEADING3:
            add_styled_heading3(doc, _clean_inline(str(content)))
        elif btype == BlockType.PARA:
            if re.match(r"^Figure\s+[\w\d]+\.\d+\s*:", str(content)):
                continue
            add_body_para(doc, _clean_inline(str(content)))
        elif btype == BlockType.BULLET:
            raw = str(content)
            level = 1 if raw.startswith("  ") else 0
            add_bullet_para(doc, _clean_inline(raw.strip()), level=level)
        elif btype == BlockType.NUMBERED:
            add_numbered_para(doc, _clean_inline(str(content)))
        elif btype == BlockType.CODE:
            add_code_block(doc, str(content))
        elif btype == BlockType.TABLE:
            headers, rows = _parse_table_rows(list(content))  # type: ignore[arg-type]
            if headers:
                table_counter[0] += 1
                ch = chapter_num or "1"
                add_styled_table(doc, headers, rows,
                                 caption=f"Table {ch}.{table_counter[0]}")
        elif btype == BlockType.FIGURE:
            alt_text, md_src = content  # type: ignore[misc]
            png = _resolve_png(str(md_src))
            if png:
                cap = FIGURE_CAPTIONS.get(png.name, f"Figure: {_clean_inline(str(alt_text))}")
                add_figure(doc, png, cap)
            else:
                add_figure_caption(doc, _clean_inline(str(alt_text)))


# ===========================================================================
# FRONT MATTER BUILDERS
# ===========================================================================

def _centered(doc, txt, size, *, bold=False, color=BRAND_MID, after=3):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, 0, after)
    return p


def build_cover_page(doc: Document) -> None:
    for _ in range(3):
        set_spacing(doc.add_paragraph(), 0, 0)

    _centered(doc, "The University of Jordan", 14, bold=True, color=BRAND_DARK, after=2)
    _centered(doc, "King Abdullah II School of Information Technology", 12, after=2)
    _centered(doc, "Department of Computer Science", 11, after=24)

    div = _centered(doc, "─" * 48, 10, color=BRAND_ACCENT, after=24)

    title_bg = doc.add_paragraph()
    set_para_shading(title_bg, BRAND_DARK)
    t_run = title_bg.add_run("  Parallax  ")
    t_run.font.name = "Times New Roman"
    t_run.font.size = Pt(28)
    t_run.font.bold = True
    t_run.font.color.rgb = BRAND_ACCENT
    title_bg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title_bg, 12, 4)

    subtitle = doc.add_paragraph()
    set_para_shading(subtitle, BRAND_DARK)
    s_run = subtitle.add_run("  A Dual-Perspective Cybersecurity Training Platform  ")
    s_run.font.name = "Times New Roman"
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = BRAND_LIGHT
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, 4, 20)

    _centered(doc, "─" * 48, 10, color=BRAND_ACCENT, after=20)

    _centered(doc, "Graduation Project - 2025 / 2026", 11, after=3)
    _centered(doc, "Submitted in partial fulfillment of the requirements", 11, after=3)
    _centered(doc, "for the Bachelor's Degree in Computer Science", 11, after=20)
    _centered(doc, "Academic Year 2025 - 2026", 12, bold=True, color=BRAND_DARK, after=6)
    add_page_break(doc)


def build_declaration(doc: Document) -> None:
    add_frontmatter_heading(doc, "Declaration")
    add_body_para(doc,
        "We declare that this graduation project report and the Parallax platform it "
        "describes are the result of our own work, except where explicit reference is made "
        "to the work of others. The platform was developed as an educational cybersecurity "
        "training environment. All offensive capabilities described in this report operate "
        "exclusively against isolated Docker scenario containers on internal-only networks; "
        "no real or third-party systems are targeted at any point.")
    add_body_para(doc,
        "This report has not been submitted, in whole or in part, for any other degree or "
        "qualification. Sources of information and external standards are acknowledged in "
        "the References section, and all third-party libraries and frameworks are cited "
        "where they support a technical decision.")
    for label in ("Student name and signature: ______________________",
                  "Supervisor name and signature: ______________________",
                  "Date: ______________________"):
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        set_spacing(p, 18, 0)
    add_page_break(doc)


def build_acknowledgments(doc: Document) -> None:
    add_frontmatter_heading(doc, "Acknowledgments")
    add_body_para(doc,
        "We would like to express our sincere gratitude to our supervisor and to the "
        "faculty of the King Abdullah II School of Information Technology at the University "
        "of Jordan for their guidance and support throughout this graduation project.")
    add_body_para(doc,
        "We also thank the open-source communities behind the technologies that made "
        "Parallax possible, including FastAPI, React, Vite, PostgreSQL, Redis, "
        "Elasticsearch, Docker, and xterm.js, as well as the security education community "
        "whose frameworks - OWASP, MITRE ATT&CK, NIST CSF, and PTES - shaped the learning "
        "design of the platform.")
    add_body_para(doc,
        "Finally, we thank our families and colleagues for their encouragement during the "
        "development and documentation of this project.")
    add_page_break(doc)


def build_abstract(doc: Document) -> None:
    add_frontmatter_heading(doc, "Abstract")
    add_body_para(doc,
        "Parallax is a browser-based cybersecurity training platform designed for university "
        "students. It provides a dual-perspective learning environment where students practice "
        "Red Team offensive techniques against isolated Docker scenario containers and "
        "simultaneously observe the Blue Team defensive telemetry generated by their actions. "
        "The platform integrates a Kali-style PTY terminal, a real-time SIEM event feed, "
        "structured methodology gating, a Socratic AI tutor, mission scoring, and post-session "
        "debrief reports. Three training scenarios are included in the MVP scope: SC-01 NovaMed "
        "(web application penetration testing), SC-02 Nexora (Active Directory compromise), and "
        "SC-03 Orion (phishing campaign simulation and SOC response). The entire stack is "
        "packaged as a single-node Docker Compose deployment, making it portable for university "
        "labs and graduation demonstrations. Instructor analytics provide session monitoring, "
        "student metrics, and report export capabilities.")
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.font.bold = True
    r.font.name = "Times New Roman"
    r2 = p.add_run(
        "cybersecurity training, penetration testing, SOC analysis, Docker, "
        "xterm.js, FastAPI, React, Elasticsearch, Socratic AI, dual-perspective")
    r2.font.name = "Times New Roman"
    r2.font.italic = True
    set_spacing(p, 12, 6)
    add_page_break(doc)


def build_toc_page(doc: Document) -> None:
    add_frontmatter_heading(doc, "Table of Contents")
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    add_page_break(doc)


def build_lof_page(doc: Document) -> None:
    add_frontmatter_heading(doc, "List of Figures")
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Figure"')
    add_page_break(doc)


def build_lot_page(doc: Document) -> None:
    add_frontmatter_heading(doc, "List of Tables")
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Table"')
    add_page_break(doc)


def build_abbreviations_page(doc: Document) -> None:
    add_frontmatter_heading(doc, "List of Abbreviations")
    add_styled_table(doc, ["Abbreviation", "Meaning"],
                     [[a, b] for a, b in ABBREVIATIONS])
    add_page_break(doc)


# ===========================================================================
# REFERENCES + APPENDICES
# ===========================================================================

def build_references_section(doc: Document) -> None:
    add_page_break(doc)
    add_styled_heading1(doc, "REFERENCES", None)
    if not REFERENCES_MD.exists():
        add_body_para(doc, "[References file not found]")
        return
    text = REFERENCES_MD.read_text(encoding="utf-8")
    ref_num = 0
    in_qa = False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            in_qa = "QA" in stripped or "Notes" in stripped
            continue
        if not stripped or stripped.startswith("#") or in_qa:
            continue
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            ref_num += 1
            p = doc.add_paragraph()
            run = p.add_run(f"[{ref_num}]  {_clean_inline(m.group(1))}")
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            set_spacing(p, before_pt=2, after_pt=4)


def build_appendices(doc: Document) -> None:
    add_page_break(doc)
    add_styled_heading1(doc, "APPENDICES", None)
    add_body_para(doc,
        "The appendices below contain the complete technical documentation that supports "
        "the main chapters: the requirements traceability matrix, the API and database "
        "references, the architecture atlas, the security and safety case, the scenario "
        "design dossier, the testing evidence, the deployment and operations manual, the "
        "user and maintainer manuals, accessibility notes, and the known limitations and "
        "future work register.")

    for letter, title, rel in APPENDICES:
        path = FR / rel
        add_page_break(doc)
        add_styled_heading1(doc, f"APPENDIX {letter}: {title}", letter,
                            label_word="APPENDIX")
        if not path.exists():
            add_body_para(doc, f"[Source not found: {rel}]")
            print(f"  [MISSING APPENDIX SOURCE] {rel}")
            continue
        tcount = [0]
        render_markdown(doc, path, tcount, caption_label=letter, skip_first_h1=True)
        print(f"  Appendix {letter} rendered from {rel} ({tcount[0]} tables)")


# ===========================================================================
# WORD COM -> PDF
# ===========================================================================

def convert_to_pdf_via_com(docx_path: Path, pdf_path: Path) -> bool:
    ps_script = f"""
Add-Type -AssemblyName Microsoft.Office.Interop.Word
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $doc = $word.Documents.Open("{str(docx_path).replace(chr(92), chr(92)*2)}")
    $doc.Fields.Update()
    $toc = $doc.TablesOfContents
    if ($toc.Count -gt 0) {{ for ($i=1; $i -le $toc.Count; $i++) {{ $toc.Item($i).Update() }} }}
    $tof = $doc.TablesOfFigures
    if ($tof.Count -gt 0) {{ for ($i=1; $i -le $tof.Count; $i++) {{ $tof.Item($i).Update() }} }}
    $doc.SaveAs2("{str(docx_path).replace(chr(92), chr(92)*2)}")
    $doc.SaveAs2("{str(pdf_path).replace(chr(92), chr(92)*2)}", 17)
    $doc.Close($false)
    $word.Quit()
    Write-Output "PDF_OK"
}} catch {{
    Write-Output "PDF_FAIL: $_"
}}
"""
    tmp_ps = Path(os.environ.get("TEMP", ".")) / "parallax_pdf_v3.ps1"
    tmp_ps.write_text(ps_script, encoding="utf-8")
    result = subprocess.run(
        ["powershell", "-ExecutionPolicy", "Bypass", "-File", str(tmp_ps)],
        capture_output=True, text=True, timeout=240)
    output = result.stdout.strip()
    print(f"  Word COM output: {output}")
    return "PDF_OK" in output


# ===========================================================================
# MAIN
# ===========================================================================

def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print("Parallax Report Compiler v3 - Final Edition")
    print("=" * 55)

    doc = Document()
    _configure_document(doc)
    _configure_heading_styles(doc)
    add_page_number_footer(doc)

    print("  Front matter: cover, declaration, acknowledgments, abstract...")
    build_cover_page(doc)
    build_declaration(doc)
    build_acknowledgments(doc)
    build_abstract(doc)
    print("  Front matter: TOC / LOF / LOT / abbreviations...")
    build_toc_page(doc)
    build_lof_page(doc)
    build_lot_page(doc)
    build_abbreviations_page(doc)

    set_page_number_restart(doc, fmt="decimal", start=1)

    table_counter = [0]
    for i, fname in enumerate(CHAPTER_FILES):
        md_path = CHAPTERS_DIR / fname
        if not md_path.exists():
            print(f"  [MISSING] {fname}")
            continue
        table_counter[0] = 0
        print(f"  Rendering {fname}...")
        render_chapter(doc, md_path, table_counter)
        if i < len(CHAPTER_FILES) - 1:
            add_page_break(doc)

    print("  Building references...")
    build_references_section(doc)
    print("  Building appendices...")
    build_appendices(doc)

    print(f"  Saving DOCX -> {OUT_DOCX.name}")
    doc.save(str(OUT_DOCX))
    print(f"  [OK] DOCX saved: {OUT_DOCX.stat().st_size:,} bytes")

    print("  Converting to PDF via Word COM (updates TOC/LOF/LOT)...")
    if convert_to_pdf_via_com(OUT_DOCX, OUT_PDF):
        print(f"  [OK] PDF saved:  {OUT_PDF.stat().st_size:,} bytes")
    else:
        print("  [WARN] Word COM conversion failed - DOCX is still good.")
        print("         Open the DOCX and run References > Update Table to populate TOC/LOF/LOT.")

    print("\n[DONE] Compilation complete.")
    print(f"   DOCX: {OUT_DOCX}")
    print(f"   PDF:  {OUT_PDF}")


if __name__ == "__main__":
    main()
